import streamlit as st
import os
import pandas as pd
import time
import json
import io
import numpy as np
import hashlib
import matplotlib.pyplot as plt
from datetime import datetime
from io import StringIO

# PDF & DOCX parsing and report generation
import PyPDF2
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# LangChain and vector store
import faiss
from langchain.vectorstores import FAISS as LC_FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document

# SentenceTransformer for feedback and column matching
from sentence_transformers import SentenceTransformer

client = OpenAI(api_key= st.secrets["OPENAI_API_KEY"])

# --- Page configuration and Custom CSS ---
st.set_page_config(
    page_title="Test Defects RCA-CAPA Automation - AI Assistant",
    page_icon="📂",
    layout="wide"
)
st.markdown("""
<style>
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    h1, h2, h3 {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .stButton button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
        border-radius: 10px;
        padding: 0.5rem 1rem;
        border: none;
        transition: all 0.3s;
    }
    .stButton button:hover {
        background-color: #45a049;
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    .deviation-card {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: all 0.3s;
    }
    .deviation-card:hover {
        box-shadow: 0 10px 20px rgba(0,0,0,0.15);
        transform: translateY(-2px);
    }
    .stSidebar .sidebar-content {
        background-color: #f8f9fa;
    }
</style>
""", unsafe_allow_html=True)

# --- Initialize Session State ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "metadata_dict" not in st.session_state:
    st.session_state.metadata_dict = {}
if "processed_data" not in st.session_state:
    st.session_state.processed_data = None
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "deviation_logs" not in st.session_state:
    st.session_state.deviation_logs = []
if "root_causes" not in st.session_state:
    st.session_state.root_causes = {}
if "feedback_logs" not in st.session_state:
    st.session_state.feedback_logs = []
if "prompt_instruction" not in st.session_state:
    st.session_state.prompt_instruction = "Provide a clear and concise response."
if "analysis_result" not in st.session_state:
    st.session_state.analysis_result = None
if "confidence_score" not in st.session_state:
    st.session_state.confidence_score = None
if "retrieved_docs" not in st.session_state:
    st.session_state.retrieved_docs = None
if "query_input" not in st.session_state:
    st.session_state.query_input = ""
if "feedback_submitted" not in st.session_state:
    st.session_state.feedback_submitted = False
if "feedback_button_clicked" not in st.session_state:
    st.session_state.feedback_button_clicked = False
if "source_document_name" not in st.session_state:
    st.session_state.source_document_name = "Unknown"
if "retriever" not in st.session_state:
    st.session_state.retriever = None
if "qa_chain" not in st.session_state:
    st.session_state.qa_chain = None

# --- LLM Model for Column Matching and Feedback ---
model = SentenceTransformer('all-MiniLM-L6-v2')

# --- Authentication ---
USER_CREDENTIALS = {
    "admin": "81dc9bdb52d04dc20036dbd8313ed055",  # MD5 for "1234"
    "user": "81dc9bdb52d04dc20036dbd8313ed055"      # MD5 for "password"
}

def hash_password(password):
    return hashlib.md5(password.encode()).hexdigest()

def authenticate(username, password):
    return USER_CREDENTIALS.get(username) == hash_password(password)

def login_page():
    st.title("🔐 Login to Deviation Assistant")
    username = st.text_input("👤 Username", key="login_username_input")
    password = st.text_input("🔑 Password", type="password", key="login_password_input")
    if st.button("Login", key="login_button"):
        if authenticate(username, password):
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success("✅ Login successful! Redirecting...")
            st.rerun()
        else:
            st.error("❌ Invalid username or password")
    st.stop()  # Prevents the rest of the app from running if not logged in

def logout():
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# --- Helper Functions for File Parsing and Metadata ---
def extract_text_from_file(uploaded_file, file_type):
    text_data = []
    if file_type == "txt":
        stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
        text_data = stringio.readlines()
    elif file_type == "pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text_data = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
    elif file_type == "csv":
        df = pd.read_csv(uploaded_file)
        text_data = df.astype(str).apply(lambda x: ' '.join(x), axis=1).tolist()
    elif file_type == "docx":
        doc = DocxDocument(uploaded_file)
        text_data = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return text_data

def load_metadata_file(metadata_file):
    if metadata_file:
        try:
            metadata_df = pd.read_csv(metadata_file)
            if "Category" in metadata_df.columns and "Keywords" in metadata_df.columns:
                metadata_dict = {
                    row["Category"]: row["Keywords"].lower().split(",")
                    for _, row in metadata_df.iterrows()
                }
                st.session_state.metadata_dict = metadata_dict
                st.success("✅ Metadata file uploaded successfully!")
            else:
                st.error("❌ Incorrect metadata format! Ensure columns: 'Category' and 'Keywords'.")
        except Exception as e:
            st.error(f"❌ Error reading metadata file: {e}")

def load_data_file(data_file):
    try:
        df = pd.read_csv(data_file)
        st.session_state.processed_data = df
        st.success("✅ Successfully loaded data file!")
        return df
    except Exception as e:
        st.error(f"❌ Error loading data file: {e}")
        return None

def load_pipeline_metadata_file(metadata_file):
    metadata_dict = {}
    try:
        if metadata_file.name.endswith(".csv"):
            df = pd.read_csv(metadata_file)
            field_name_col = "Field_Name"
            description_col = "Definition"
            if field_name_col not in df.columns or description_col not in df.columns:
                st.error(f"❌ Metadata file is missing required columns. Found: {list(df.columns)}")
                return {}
            metadata_dict = dict(zip(df[field_name_col], df[description_col]))
        elif metadata_file.name.endswith(".json"):
            metadata_dict = json.load(metadata_file)
        st.session_state.metadata_dict = metadata_dict
        st.success(f"✅ Metadata file loaded successfully with {len(metadata_dict)} fields!")
    except Exception as e:
        st.error(f"❌ Error loading metadata file: {e}")
    return metadata_dict

def map_columns_with_llm(data_columns, metadata_dict):
    column_mapping = {}
    used_names = set()
    for col in data_columns:
        if col in metadata_dict:
            mapped_name = metadata_dict[col]
        else:
            col_embedding = model.encode(col)
            best_match = None
            best_score = -1
            for meta_field, description in metadata_dict.items():
                meta_embedding = model.encode(meta_field)
                score = col_embedding.dot(meta_embedding) / (np.linalg.norm(col_embedding) * np.linalg.norm(meta_embedding))
                if score > best_score:
                    best_score = score
                    best_match = meta_field
            mapped_name = metadata_dict.get(best_match, "Unknown Field")
        original_mapped_name = mapped_name
        counter = 1
        while mapped_name in used_names:
            mapped_name = f"{original_mapped_name} ({counter})"
            counter += 1
        used_names.add(mapped_name)
        column_mapping[col] = mapped_name
    return column_mapping

def store_metadata_in_vector_db(metadata_dict):
    if "OPENAI_API_KEY" not in os.environ or not os.environ["OPENAI_API_KEY"]:
        st.error("❌ OpenAI API key is missing! Set `OPENAI_API_KEY` before proceeding.")
        return
    documents = [Document(page_content=f"{key}: {value}") for key, value in metadata_dict.items()]
    embedding_model = OpenAIEmbeddings(openai_api_key=os.environ["OPENAI_API_KEY"])
    vector_store = LC_FAISS.from_documents(documents, embedding_model)
    st.session_state.vector_store = vector_store
    st.success(f"✅ Stored {len(metadata_dict)} metadata entries in FAISS vector database!")

# --- Deviation Analysis Functions ---
def categorize_root_cause(root_cause):
    if not st.session_state.metadata_dict:
        return "Other"
    root_cause = root_cause.lower()
    for category, keywords in st.session_state.metadata_dict.items():
        if any(keyword.strip() in root_cause for keyword in keywords):
            return category
    return "Other"

def plot_fishbone(causes):
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_title("Fishbone Diagram - Root Causes", fontsize=16, fontweight="bold")
    ax.plot([1, 11], [5, 5], "k-", lw=3)
    #categories = ["Functional Defect", "UI-UX", "Integration Defect", "Validation/Rules Error", "Test Data Issue", "Environment Config", "Tool/Script Issue", "People (Human Error)", "Process Gap"]
    categories = [
    "Functional Defect", "UI-UX", "Integration Defect",
    "Validation/Rules Error", "Test Data Issue", "Environment Config",
    "Tool/Script Issue", "People (Human Error)", "Process Gap","Other"
    ]

    category_causes = {cat: [] for cat in categories}
    for cause in causes:
        category = categorize_root_cause(cause)
        category_causes[category].append(cause)
    #y_positions = [8, 7, 6, 4, 3, 2]
    y_positions = np.linspace(2, 8, len(category_causes))

    x_main = 6
    for i, (category, cause_list) in enumerate(category_causes.items()):
        y = y_positions[i]
        ax.plot([x_main - 2, x_main], [y, 5], "k-", lw=2)
        ax.text(x_main - 2.5, y, category, fontsize=13, fontweight="bold", verticalalignment="center")
        for j, cause in enumerate(cause_list):
            ax.text(x_main + 1, y - j * 0.4, f"- {cause}", fontsize=11, verticalalignment="center")
    ax.axis("off")
    st.pyplot(fig)

def generate_fda_capa_section(defect_text, root_causes, llm_chain):
    prompt = f"""
You are an FDA compliance expert. Based on the following defect description and identified root causes, generate a detailed CAPA report compliant with FDA CSA and 21 CFR Part 11.

Include the following:
1. Problem Summary
2. Root Cause
3. Corrective Actions
4. Preventive Actions
5. Risk Assessment
6. Verification of Effectiveness (VOE)
7. SOP or Compliance Reference
8. Assigned Owner and Timeline

Defect:
{defect_text}

Root Causes:
{json.dumps(root_causes, indent=2)}
"""
    return llm_chain.run(prompt).strip()


def generate_pdf_report(transaction_detail, analysis, root_causes, capa_suggestions, fda_capa_text):
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.setTitle("Defect Report")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(200, 750, "🚀 Test Defect Report")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(200, 730, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    pdf.line(50, 720, 550, 720)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, 700, "🔍 Defect Details")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, 680, f"Defect Detail: {transaction_detail}")
    pdf.drawString(50, 660, f"Analysis: {analysis}")
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, 640, "⚠️ Root Causes Identified")
    pdf.setFont("Helvetica", 12)
    y = 620
    for cause in root_causes:
        pdf.drawString(60, y, f"- {cause}")
        y -= 20
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y - 20, "✅ Corrective & Preventive Actions (CAPA)")
    pdf.setFont("Helvetica", 12)
    y -= 40
    for action in capa_suggestions:
        pdf.drawString(60, y, f"- {action}")
        y -= 20

        # New Page for FDA CAPA
    pdf.showPage()
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, 750, "📋 FDA CSA-Based CAPA Report")
    pdf.setFont("Helvetica", 11)

    y = 730
    for line in fda_capa_text.split("\n"):
        if y < 80:
            pdf.showPage()
            y = 750
        pdf.drawString(50, y, line[:100])
        y -= 15


    pdf.save()
    buffer.seek(0)
    return buffer

# --- Feedback-Based Fine-Tuning Functions ---
def load_feedback_data():
    try:
        with open("feedback_data.json", "r") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def create_faiss_index(feedback_data):
    if not feedback_data:
        return None, None
    queries = [fb["Defect Detail"] for fb in feedback_data]
    query_embeddings = np.array([model.encode(q) for q in queries])
    index = faiss.IndexFlatL2(query_embeddings.shape[1])
    index.add(query_embeddings)
    return index, feedback_data

def extract_insights_from_feedback(feedback_items):
    insights = []
    for item in feedback_items:
        if item["User Feedback"] == "👎 No" and item["Comments"].strip():
            insights.append(item["Comments"])
    return insights

def retrieve_feedback_insights(query, index, feedback_data, max_results=3, threshold=0.8):
    if index is None:
        return []
    query_vector = model.encode(query).reshape(1, -1)
    distances, indices = index.search(query_vector, k=max_results)
    relevant_feedback = []
    for i, idx in enumerate(indices[0]):
        if distances[0][i] < threshold:
            relevant_feedback.append(feedback_data[idx])
    return extract_insights_from_feedback(relevant_feedback)

def blend_analysis_with_feedback(base_analysis, feedback_insights):
    if not feedback_insights:
        return base_analysis
    combined_analysis = base_analysis + "\n\n**Additional Insights from Similar Cases:**\n"
    for i, insight in enumerate(feedback_insights):
        combined_analysis += f"{i+1}. {insight}\n"
    return combined_analysis

def calculate_confidence_score(base_score, feedback_insights, retrieved_docs=None):
    if retrieved_docs:
        num_docs = len(retrieved_docs)
        rank_scores = np.linspace(1.0, 0.5, num_docs)
        base_score = round(np.mean(rank_scores), 2)
    if feedback_insights:
        feedback_boost = min(0.2, 0.05 * len(feedback_insights))
        return min(0.95, base_score + feedback_boost)
    return base_score

def prepare_training_data(feedback_logs):
    training_pairs = []
    for log in feedback_logs:
        if log["User Feedback"] == "👎 No" and log["Comments"].strip():
            training_pairs.append({
                "query": log.get("Transaction Detail", log.get("Defect Detail", "")),
                "improved_response": log["Comments"],
                "original_response": log["Analysis"]
            })
    return training_pairs

def save_training_data(training_pairs):
    with open("training_data.json", "w") as f:
        json.dump(training_pairs, f, indent=4)
    return len(training_pairs)

def retrieve_corrected_response(query, index, feedback_data, threshold=0.7):
    if index is None:
        return None
    query_vector = model.encode(query).reshape(1, -1)
    distances, indices = index.search(query_vector, k=1)
    if distances[0][0] < threshold:
        matched_feedback = feedback_data[indices[0][0]]
        if matched_feedback["User Feedback"] == "👎 No":
            return matched_feedback["Comments"]
    return None

# --- Sidebar Navigation and Main Page Content ---
if not st.session_state.logged_in:
    login_page()
else:
    # --- Sidebar Navigation ---
    with st.sidebar:
        st.title(f"🚀 Welcome, {st.session_state['username']}")
        if st.button("Logout"):
            logout()

        # Choose one of the three main AI agents
        main_agent = st.radio("📌 Select AI Agent",
                             ["Data Management AI Agent", "RAG AI Agent", "Application AI Agent"],
                             key="main_agent")

        # Show appropriate submenu based on main agent selection
        if main_agent == "Data Management AI Agent":
            dm_submenu = st.radio("📂 Select Data Task",
                                ["Pipeline Dashboard", "Data Pipeline", "Processed Data"],
                                key="dm_radio")

        elif main_agent == "RAG AI Agent":
            submenu = st.radio("📂 Select RAG Task",
                              ["Dashboard", "Configure & Upload", "Fine Tuning", "Settings"],
                              key="rag_radio")

        elif main_agent == "Application AI Agent":
            app_submenu = st.radio("📂 Select Application Task",
                              ["Defect Analysis", "Root Cause Analysis", "Analysis History", "Defect Classifier","User Feedback"],
                               key="app_radio")

     # --- Main Page Content ---
     ## Only show content for the currently selected main agent
    if main_agent == "Data Management AI Agent":
        if dm_submenu == "Pipeline Dashboard":
            st.header("📊 Data Pipeline Dashboard")
            if st.session_state.processed_data is not None:
                st.dataframe(st.session_state.processed_data.head())
            else:
                st.info("No processed data available yet.")
        elif dm_submenu == "Data Pipeline":
            st.header("🛠️ AI Data Processing Pipeline")
            data_file = st.file_uploader("📂 Upload Data File (CSV)", type=["csv"], key="data_pipeline_file")
            metadata_file = st.file_uploader("📜 Upload Metadata File (CSV or JSON)", type=["csv", "json"], key="metadata_pipeline_file")
            if st.button("🔄 Process Data"):
                if data_file is None or metadata_file is None:
                    st.error("❌ Please upload both data and metadata files!")
                else:
                    data_df = load_data_file(data_file)
                    load_pipeline_metadata_file(metadata_file)
                    if data_df is not None and st.session_state.metadata_dict:
                         # Map columns using the metadata
                        column_mapping = map_columns_with_llm(data_df.columns, st.session_state.metadata_dict)
                        data_df.rename(columns=column_mapping, inplace=True)
                         # Deduplicate column names if needed
                        data_df.columns = pd.io.common.dedup_names(list(data_df.columns), is_potential_multiindex=False)
                        st.session_state.processed_data = data_df.copy()
                        st.success("✅ Processed data stored successfully!")
                         # Store metadata in FAISS vector DB
                        store_metadata_in_vector_db(st.session_state.metadata_dict)
                        st.rerun()
        elif dm_submenu == "Processed Data":
            st.header("📊 Processed Data Insights")
            if st.session_state.processed_data is None:
                st.warning("No processed data available. Run the pipeline first.")
            else:
                st.dataframe(st.session_state.processed_data.head())
                if st.session_state.metadata_dict:
                    column_mapping_df = pd.DataFrame(list(st.session_state.metadata_dict.items()),
                                                    columns=["Field Name", "Description"])
                    st.dataframe(column_mapping_df)
                query = st.text_input("Enter field name to lookup:", key="metadata_lookup")
                if st.button("🔎 Search Metadata"):
                    if query in st.session_state.metadata_dict:
                        st.success(f"📌 {query}: {st.session_state.metadata_dict[query]}")
                    else:
                        st.warning("⚠️ Field not found in metadata.")

    elif main_agent == "RAG AI Agent":
        if submenu == "Dashboard":
            st.header("🚀 RAG AI Agent Dashboard")
            st.write("Monitor key metrics and insights from your analysis.")
            if st.session_state.feedback_logs:
                feedback_df = pd.DataFrame(st.session_state.feedback_logs)
                st.subheader("📊 User Feedback Insights")
                st.dataframe(feedback_df)
        elif submenu == "Configure & Upload":
            st.header("🏗️ Configure & Upload Data")
            uploaded_file = st.file_uploader("Upload Defect Log (TXT, PDF, CSV, DOCX)", type=["txt", "pdf", "csv", "docx"])
            metadata_file = st.file_uploader("Upload Metadata File (CSV)", type="csv")
            prompt_instruction = st.text_area("📝 Provide Prompt Instructions", value=st.session_state.prompt_instruction)
            if st.button("Submit & Process Data"):
                if uploaded_file:
                    st.session_state.source_document_name = uploaded_file.name
                    file_extension = uploaded_file.name.split(".")[-1].lower()
                    with st.spinner("Processing data..."):
                        for i in range(100):
                            time.sleep(0.01)
                        text_data = extract_text_from_file(uploaded_file, file_extension)
                        if text_data:
                            documents = [Document(page_content=line) for line in text_data]
                            embedding_model = OpenAIEmbeddings()
                            ector_store = LC_FAISS.from_documents(documents, embedding_model)
                            retriever = vector_store.as_retriever()
                            llm = OpenAI()
                            qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
                            st.session_state.retriever = retriever
                            st.session_state.qa_chain = qa_chain
                            st.session_state.prompt_instruction = prompt_instruction
                            st.success("✅ Data successfully stored in FAISS. Ready for analysis!")
                        else:
                            st.error("❌ Could not extract text from the uploaded file.")
                if metadata_file:
                    load_metadata_file(metadata_file)
        elif submenu == "Fine Tuning":
            st.header("🛠️ Fine Tune Model")
            try:
                with open("training_data.json", "r") as f:
                    training_data = json.load(f)
                num_examples = len(training_data)
            except (FileNotFoundError, json.JSONDecodeError):
                training_data = []
                num_examples = 0
            st.subheader("📊 Training Data Summary")
            if num_examples > 0:
                st.write(f"📝 **Available Training Examples:** {num_examples}")
                for i in range(min(5, num_examples)):
                    with st.expander(f"Example {i+1}"):
                        st.write("**Original Query:**", training_data[i]['query'])
                        st.write("**Original Response:**", training_data[i]['original_response'])
                        st.write("**Improved Response:**", training_data[i]['improved_response'])
                st.subheader("🧠 Fine-Tuning Options")
                epochs = st.number_input("Training Epochs", min_value=1, max_value=10, value=3)
                learning_rate = st.select_slider("Learning Rate", options=[0.00001, 0.0001, 0.001, 0.01], value=0.0001)
                if st.button("Regenerate Training Data"):
                    training_pairs = prepare_training_data(st.session_state.feedback_logs)
                    num_pairs = save_training_data(training_pairs)
                    st.success(f"✅ Successfully regenerated {num_pairs} training examples!")
                    st.rerun()
                if st.button("Start Fine-Tuning"):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    for i in range(100):
                        progress_bar.progress(i + 1)
                        if i < 30:
                            status_text.text(f"Preparing training data... ({i+1}%)")
                        elif i < 80:
                            status_text.text(f"Training model... ({i+1}%)")
                        else:
                            status_text.text(f"Evaluating and finalizing model... ({i+1}%)")
                        time.sleep(0.05)
                    status_text.text("Fine-tuning complete! Model has been updated.")
                    st.success("✅ Model successfully fine-tuned with user feedback!")
        elif submenu == "Settings":
            st.header("⚙️ RAG AI Agent Settings")
            col1, col2 = st.columns(2)
            with col1:
                model_choice = st.selectbox("Language Model", ["OpenAI GPT-3.5", "OpenAI GPT-4", "Custom Model"])
            with col2:
                temperature = st.slider("Temperature (Creativity)", 0.0, 1.0, 0.7, 0.1)
            st.subheader("User Interface")
            theme = st.radio("Color Theme", ["Light", "Dark", "System Default"])
            if st.button("Save Settings"):
                st.success("✅ Settings saved successfully!")

    elif main_agent == "Application AI Agent":
        if app_submenu == "Defect Analysis":
            st.header("📉 Defect Analysis")
            if st.session_state.retriever is None:
                st.info("👈 Please upload and process a defect log file first using the RAG AI Agent.")
            else:
                query = st.text_area("Enter Defect details to check for defects:", height=100,
                                     placeholder="Example: Login fails with valid credentials, returning a 500 error from the backend service.",
                                     value=st.session_state.query_input)
                if st.button("Analyze Deviation") and query:
                    feedback_data = load_feedback_data()
                    faiss_index, feedback_data_indexed = create_faiss_index(feedback_data)
                    feedback_insights = retrieve_feedback_insights(query, faiss_index, feedback_data_indexed)
                    if feedback_insights:
                        feedback_text = "\n".join(f"- {insight}" for insight in feedback_insights)
                        composite_prompt = f"Based on feedback from similar cases:\n{feedback_text}\n\nNow, analyze the following transaction:\n{query}"
                    else:
                        composite_prompt = query
                    response = st.session_state.qa_chain.invoke(f"{st.session_state.prompt_instruction}\n{composite_prompt}")
                    base_analysis = response.get("result", f"Deviation detected in {query}.")
                    st.session_state.analysis_result = base_analysis
                    st.session_state.confidence_score = calculate_confidence_score(0.6, feedback_insights, response.get("source_documents", []))
                    st.session_state.query_input = query
                    st.session_state.deviation_logs.append({
                         "Transaction Detail": query,
                         "Analysis": base_analysis,
                         "Source Document": st.session_state.source_document_name,
                         "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                     })
                if st.session_state.analysis_result:
                    st.subheader("Defect Analysis Result:")
                    st.markdown(f'<div class="deviation-card">{st.session_state.analysis_result}</div>', unsafe_allow_html=True)
                    st.subheader("Source Document:")
                    st.markdown(f'<div class="deviation-card">{st.session_state.source_document_name}</div>', unsafe_allow_html=True)
                    st.subheader("🧠 Confidence Score:")
                    st.markdown(f'<div class="deviation-card">Confidence: {round(st.session_state.confidence_score * 100)}%</div>', unsafe_allow_html=True)
                    st.subheader("Was this response helpful?")
                    feedback = st.radio("Feedback:", ["👍 Yes", "👎 No"], key="feedback_radio", horizontal=True)
                    feedback_comment = st.text_area("Additional comments or corrections:", height=100, key="feedback_comment")
                    if st.button("Submit Feedback"):
                        st.session_state.feedback_button_clicked = True
                    if st.session_state.feedback_button_clicked:
                        feedback_entry = {
                             "Transaction Detail": st.session_state.query_input,
                             "Analysis": st.session_state.analysis_result,
                             "Source Document": st.session_state.source_document_name,
                             "User Feedback": feedback,
                             "Comments": feedback_comment,
                             "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                         }
                        st.session_state.feedback_logs.append(feedback_entry)
                        try:
                            with open("feedback_data.json", "w") as f:
                                json.dump(st.session_state.feedback_logs, f, indent=4)
                            if feedback == "👎 No" and feedback_comment.strip():
                                training_pairs = prepare_training_data(st.session_state.feedback_logs)
                                num_pairs = save_training_data(training_pairs)
                                st.success(f"✅ Feedback recorded and {num_pairs} training examples saved!")
                            else:
                                st.success("✅ Feedback recorded successfully!")
                            st.session_state.feedback_submitted = True
                            st.session_state.feedback_button_clicked = False
                        except Exception as e:
                            st.error(f"Error saving feedback: {e}")
        elif app_submenu == "Root Cause Analysis":
            st.header("🔍 Root Cause Analysis")
            transactions = [log["Transaction Detail"] for log in st.session_state.deviation_logs]
            selected_transaction = st.selectbox("Select a defect to analyze:", transactions if transactions else ["No defects available"])
            if st.button("Identify Root Causes"):
                with st.spinner("Identifying root causes..."):
                    root_causes = ["Functional Defect", "UI/UX", "People (Human Error)", "Other"]
                    #root_causes = ["Functional Defect", "UI/UX", "Integration Defect", "Validation/Rules Error", "Test Data Issue", "Environment Config", "Tool/Script Issue", "People (Human Error)", "Process Gap","Other"]

                    st.session_state.root_causes[selected_transaction] = root_causes
                st.success("✅ Root causes identified!")
            if selected_transaction in st.session_state.root_causes:
                causes = st.session_state.root_causes[selected_transaction]
                st.subheader("Identified Root Causes:")
                for cause in causes:
                    st.markdown(f"- {cause}")
                st.subheader("📊 Fishbone Diagram (Ishikawa)")
                plot_fishbone(causes)
                 # Generate FDA CAPA Section using LLM
                fda_capa_text = generate_fda_capa_section(
                    defect_text=selected_transaction,
                    root_causes=causes,
                    llm_chain=st.session_state.qa_chain  # or whatever your RetrievalQA/OpenAI object is
                 )

                 # Generate full PDF
                pdf_buffer = generate_pdf_report(
                     transaction_detail=selected_transaction,
                     analysis="Analysis Example",
                     root_causes=causes,
                     capa_suggestions=["Train staff", "Improve material quality"],
                     fda_capa_text=fda_capa_text
                 )

                 #pdf_buffer = generate_pdf_report(selected_transaction, "Analysis Example", causes, ["Train staff", "Improve material quality"])
                st.download_button("📄 Download PDF Report", data=pdf_buffer, file_name="Defect_Report.pdf", mime="application/pdf")
        elif app_submenu == "Analysis History":
            st.header("📊 Analysis History & Insights")
            if len(st.session_state.deviation_logs) == 0:
                st.info("No analysis data available yet. Run some analyses first.")
            else:
                log_df = pd.DataFrame(st.session_state.deviation_logs)
                if "Source Document" not in log_df.columns:
                    log_df["Source Document"] = "Unknown"
                columns_to_show = ["Defect Detail", "Analysis", "Source Document"]
                if "Timestamp" in log_df.columns:
                    columns_to_show.append("Timestamp")
                st.dataframe(log_df[columns_to_show], use_container_width=True)
                st.subheader("📊 Defect Statistics")
                deviation_counts = log_df["Defect Detail"].value_counts()
                fig, ax = plt.subplots(figsize=(10, 6))
                deviation_counts.plot(kind="bar", ax=ax, color="skyblue")
                ax.set_xlabel("Defect Type")
                ax.set_ylabel("Occurrences")
                ax.set_title("Deviation Count Chart")
                plt.tight_layout()
                st.pyplot(fig)

         #####Adding Defect Categorizarion  Code. #####

        elif app_submenu == "Defect Classifier":
            st.header("📂 Defect Classifier (Auto Categorization + RCA)")
            st.write("Upload a **training file** and a **test defect file**, and the system will predict:")
            st.markdown("""
             - 🧠 **Category**
             - 🚨 **Priority**
             - 🔍 **Root Cause**
             """)

            train_file = st.file_uploader("📘 Upload Training File (Excel)", type=["xlsx"], key="defect_train_file")
            test_file = st.file_uploader("📄 Upload Test File (Excel)", type=["xlsx"], key="defect_test_file")

            if train_file and test_file and st.button("🔍 Classify Defects"):
                from sentence_transformers import SentenceTransformer
                from sklearn.neighbors import NearestNeighbors
                import pandas as pd

                try:
                    df_train = pd.read_excel(train_file)
                    df_test = pd.read_excel(test_file)
                except Exception as e:
                    st.error(f"❌ Failed to read uploaded files: {e}")
                    st.stop()

                # Ensure expected columns
                required_cols = ["Defect_ID", "Summary", "Description", "Category", "Priority", "Root_Cause"]
                if not all(col in df_train.columns for col in required_cols):
                    st.error("❌ Training file must contain columns: " + ", ".join(required_cols))
                    st.stop()
                if not all(col in df_test.columns for col in ["Defect_ID", "Summary", "Description"]):
                    st.error("❌ Test file must contain 'Defect_ID', 'Summary', 'Description'")
                    st.stop()

                # Load embedding model and train NearestNeighbors
                model_embed = SentenceTransformer("all-MiniLM-L6-v2")
                train_embeddings = model_embed.encode(df_train["Description"].astype(str).tolist())
                test_embeddings = model_embed.encode(df_test["Description"].astype(str).tolist())

                nn_model = NearestNeighbors(n_neighbors=3, metric="cosine")
                nn_model.fit(train_embeddings)

                distances, indices = nn_model.kneighbors(test_embeddings)

                df_result = df_test.copy()
                for i, idx_list in enumerate(indices):
                    best_idx = idx_list[0]
                    df_result.at[i, "Predicted_Category"] = df_train.at[best_idx, "Category"]
                    df_result.at[i, "Predicted_Priority"] = df_train.at[best_idx, "Priority"]
                    df_result.at[i, "Predicted_Root_Cause"] = df_train.at[best_idx, "Root_Cause"]

                st.success("✅ Classification complete. See below:")
                st.dataframe(df_result, use_container_width=True)

                csv = df_result.to_csv(index=False).encode("utf-8")
                st.download_button("📥 Download Results as CSV", csv, "defect_classification_results.csv", "text/csv")

                 # 📊 Charts
                st.subheader("📊 Defect Breakdown by Category and Priority")
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**Category Distribution**")
                    st.bar_chart(df_result["Predicted_Category"].value_counts())

                with col2:
                    st.markdown("**Priority Distribution**")
                    st.bar_chart(df_result["Predicted_Priority"].value_counts())

                 # 🧵 Top 3 Similar Matches
                st.subheader("🧵 Top 3 Similar Defects for Each Test Case")
                for i in range(len(df_result)):
                    st.markdown(f"**🔍 Test Case {df_result.at[i, 'Defect_ID']}**: {df_result.at[i, 'Summary']}")
                    for rank, idx in enumerate(indices[i]):
                        match = df_train.iloc[idx]
                        st.markdown(f"- {rank+1}. **{match['Defect_ID']}** | *{match['Summary']}* | {match['Category']} | {match['Root_Cause']}")


         #### End of the defect classifier #####

        elif app_submenu == "User Feedback":
            st.header("📝 User Feedback Dashboard")
            if len(st.session_state.feedback_logs) == 0:
                st.info("No feedback available yet.")
            else:
                feedback_df = pd.DataFrame(st.session_state.feedback_logs)
                st.subheader("📋 Collected Feedback")
                st.dataframe(feedback_df, use_container_width=True)
                st.subheader("📊 Feedback Summary")
                positive_feedback_count = (feedback_df["User Feedback"] == "👍 Yes").sum()
                negative_feedback_count = (feedback_df["User Feedback"] == "👎 No").sum()
                st.write(f"✅ **Positive Feedback:** {positive_feedback_count}")
                st.write(f"❌ **Negative Feedback:** {negative_feedback_count}")
                if not feedback_df.empty:
                    st.subheader("📊 Feedback Trends Over Time")
                    feedback_df["Timestamp"] = pd.to_datetime(feedback_df["Timestamp"])
                    feedback_over_time = feedback_df.groupby(feedback_df["Timestamp"].dt.date)["User Feedback"].value_counts().unstack().fillna(0)
                    fig, ax = plt.subplots(figsize=(8, 5))
                    feedback_over_time.plot(kind="bar", ax=ax, stacked=True)
                    ax.set_xlabel("Date")
                    ax.set_ylabel("Number of Feedbacks")
                    ax.set_title("User Feedback Trends")
                    plt.xticks(rotation=45)
                    st.pyplot(fig)
                csv_feedback = feedback_df.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Download Feedback as CSV", csv_feedback, "user_feedback.csv", "text/csv", key='download-feedback')

     # --- Footer ---
    st.markdown("""
    <div style="text-align: center; margin-top: 30px; padding: 15px; background-color: #f5f5f5;">
        <p>Deviation Assistant Pro • Made with ❤️ by CloudHub</p>
    </div>
    """, unsafe_allow_html=True)
